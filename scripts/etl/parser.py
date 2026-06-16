"""
Attachment parser: extract plain text from uploaded files.

Supported:  .txt and text-based .pdf  (requires pdfplumber)
Unsupported: image files, scanned PDFs — returns a placeholder string instead
             of raising an exception, so batch processing is never interrupted.
"""
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

_UNSUPPORTED_PLACEHOLDER = "[IMAGE_UNSUPPORTED: 当前版本暂不支持图片解析]"
_CORRUPTED_PLACEHOLDER   = "[FILE_CORRUPTED: 文件损坏或格式错误]"
_IMAGE_SUFFIXES = frozenset({
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif",
})
_DOCX_SUFFIXES = frozenset({".doc", ".docx"})
_PDF_MIN_TEXT_CHARS = 20
_MAX_TEXT_CHARS     = 10000
_TRUNCATION_SUFFIX  = "...[文本过长已截断]"

# ── Public API ────────────────────────────────────────────────────────────────

def parse_file(path: "str | Path") -> str:
    """Extract text from a local attachment file.  Never raises."""
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in _IMAGE_SUFFIXES:
        logger.warning("\033[33m[PARSER-WARN]\033[0m Skipping image attachment: %s", path.name)
        return _UNSUPPORTED_PLACEHOLDER

    if suffix == ".pdf":
        return _parse_pdf_path(path)

    if suffix in _DOCX_SUFFIXES:
        return _parse_docx_path(path)

    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return _truncate(text)
    except Exception as exc:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m Could not read %s: %s", path.name, exc
        )
        return _CORRUPTED_PLACEHOLDER


def parse_bytes(raw: bytes, filename: str) -> str:
    """Extract text from in-memory bytes.  Never raises.

    Convenience wrapper for Netlify API workflows where attachments
    are downloaded before parsing.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in _IMAGE_SUFFIXES:
        logger.warning("\033[33m[PARSER-WARN]\033[0m Skipping image attachment: %s", filename)
        return _UNSUPPORTED_PLACEHOLDER

    if suffix == ".pdf":
        return _parse_pdf_bytes(raw, filename)

    if suffix in _DOCX_SUFFIXES:
        return _parse_docx_bytes(raw, filename)

    try:
        text = raw.decode("utf-8", errors="replace")
        return _truncate(text)
    except Exception as exc:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m Could not decode %s: %s", filename, exc
        )
        return _CORRUPTED_PLACEHOLDER

# ── Helpers ───────────────────────────────────────────────────────────────────

def _truncate(text: str) -> str:
    if len(text) > _MAX_TEXT_CHARS:
        return text[:_MAX_TEXT_CHARS] + _TRUNCATION_SUFFIX
    return text


def _parse_pdf_path(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m pdfplumber not installed; skipping PDF: %s", path.name
        )
        return "[PDF_UNSUPPORTED: 请运行 pip install pdfplumber]"

    try:
        with pdfplumber.open(path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
    except Exception as exc:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m Failed to parse PDF %s: %s", path.name, exc
        )
        return _CORRUPTED_PLACEHOLDER

    if len(text) < _PDF_MIN_TEXT_CHARS:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m PDF appears to be scanned "
            "(%d chars extracted, threshold=%d): %s",
            len(text), _PDF_MIN_TEXT_CHARS, path.name,
        )
        return _UNSUPPORTED_PLACEHOLDER

    return _truncate(text)


def _parse_pdf_bytes(raw: bytes, filename: str) -> str:
    import io

    try:
        import pdfplumber
    except ImportError:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m pdfplumber not installed; skipping PDF: %s", filename
        )
        return "[PDF_UNSUPPORTED: 请运行 pip install pdfplumber]"

    try:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
    except Exception as exc:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m Failed to parse PDF %s: %s", filename, exc
        )
        return _CORRUPTED_PLACEHOLDER

    if len(text) < _PDF_MIN_TEXT_CHARS:
        logger.warning(
            "\033[33m[PARSER-WARN]\033[0m PDF appears to be scanned "
            "(%d chars extracted, threshold=%d): %s",
            len(text), _PDF_MIN_TEXT_CHARS, filename,
        )
        return _UNSUPPORTED_PLACEHOLDER

    return _truncate(text)


def _extract_docx_text(doc) -> str:
    """Extract all text from a python-docx Document, covering paragraphs and tables.

    Word documents frequently store body content in tables rather than top-level
    paragraphs (common for formatted guides and structured reports). Reading only
    doc.paragraphs misses that content entirely, returning an empty string that
    triggers the CORRUPTED_PLACEHOLDER path.  This helper reads both sources and
    deduplicates merged table cells via their underlying XML element identity.
    """
    parts: list[str] = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    seen_cells: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def _parse_docx_path(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        logger.warning("\033[33m[PARSER-WARN]\033[0m python-docx not installed; skipping DOCX: %s", path.name)
        return "[DOCX_UNSUPPORTED: 请运行 pip install python-docx]"
    try:
        doc = Document(path)
        text = _extract_docx_text(doc)
        return _truncate(text) if text.strip() else _CORRUPTED_PLACEHOLDER
    except Exception as exc:
        logger.warning("\033[33m[PARSER-WARN]\033[0m Failed to parse DOCX %s: %s", path.name, exc)
        return _CORRUPTED_PLACEHOLDER


def _parse_docx_bytes(raw: bytes, filename: str) -> str:
    import io
    try:
        from docx import Document
    except ImportError:
        logger.warning("\033[33m[PARSER-WARN]\033[0m python-docx not installed; skipping DOCX: %s", filename)
        return "[DOCX_UNSUPPORTED: 请运行 pip install python-docx]"
    try:
        doc = Document(io.BytesIO(raw))
        text = _extract_docx_text(doc)
        return _truncate(text) if text.strip() else _CORRUPTED_PLACEHOLDER
    except Exception as exc:
        logger.warning("\033[33m[PARSER-WARN]\033[0m Failed to parse DOCX %s: %s", filename, exc)
        return _CORRUPTED_PLACEHOLDER
